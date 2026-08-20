#!/usr/bin/env python3
"""Build the 25-minute Databricks speaker script as a Word document."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x08, 0x2B, 0x46)
GOLD = RGBColor(0x8C, 0x6A, 0x22)
INK = RGBColor(0x03, 0x15, 0x22)
MUTED = RGBColor(0x5B, 0x6B, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
PALE = "F5EDD9"
NAVY_HEX = "082B46"
ROW_ALT = "E7EFF4"

OUT = Path(__file__).resolve().parents[1] / "ONR_ITSS_Databricks_Speaker_Script.docx"


def _shade(cell, hex_color: str) -> None:
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_run(run, *, bold=False, italic=False, size=11, color=INK, font="Calibri"):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)


def p(doc, text="", *, bold=False, italic=False, size=11, color=INK, space_after=8, space_before=0, align=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = 1.15
    if align:
        para.alignment = align
    if text:
        r = para.add_run(text)
        _set_run(r, bold=bold, italic=italic, size=size, color=color)
    return para


def mixed(doc, parts, *, space_after=8, space_before=0):
    """parts: list of (text, kwargs)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = 1.15
    for text, kw in parts:
        r = para.add_run(text)
        _set_run(r, **kw)
    return para


def heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    para.paragraph_format.space_after = Pt(6)
    r = para.add_run(text)
    _set_run(r, bold=True, size=16 if level == 1 else 13, color=NAVY if level == 1 else GOLD)
    # bottom border for h1
    if level == 1:
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "8C6A22")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return para


def do_this(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.left_indent = Inches(0.15)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE)
    shd.set(qn("w:val"), "clear")
    para._p.get_or_add_pPr().append(shd)
    r1 = para.add_run("[DO THIS]  ")
    _set_run(r1, bold=True, italic=True, size=11, color=GOLD)
    r2 = para.add_run(text)
    _set_run(r2, italic=True, size=11, color=INK)
    return para


def speak(doc, text):
    return p(doc, text, size=11, color=INK, space_after=8)


def table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        r = para.add_run(h)
        _set_run(r, bold=True, size=10, color=WHITE)
        _shade(cell, NAVY_HEX)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = ""
            para = cell.paragraphs[0]
            r = para.add_run(str(val))
            _set_run(r, bold=(ci == 0), size=10, color=INK)
            if ri % 2 == 1:
                _shade(cell, ROW_ALT)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return tbl


def set_header_footer(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    r = hp.add_run("UNCLASSIFIED  //  MOCK DATA  ·  ONR ITSS Factor 3  ·  Databricks / Project Compass")
    _set_run(r, bold=True, size=9, color=GOLD)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Databricks speaker script  ·  25 minutes  ·  one Key Personnel  ·  page ")
    _set_run(r, size=9, color=MUTED)
    # PAGE field
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r2 = fp.add_run()
    r2._r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    r3 = fp.add_run()
    r3._r.append(instr)
    _set_run(r3, size=9, color=MUTED)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r4 = fp.add_run()
    r4._r.append(fld2)


def build():
    doc = Document()
    set_header_footer(doc)

    p(doc, "ONR ITSS  ·  Factor 3 Technical Demonstration", bold=True, size=11, color=GOLD, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "Databricks Speaker Script", bold=True, size=26, color=NAVY, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "Project Compass  ·  Code 08 portfolio  ·  25 minutes or less", italic=True, size=12, color=MUTED, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "One Key Personnel  ·  live cloud + live repo  ·  no slides  ·  no overlays", size=11, color=INK, space_after=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    heading(doc, "How to use this script")
    speak(doc, "This is the Databricks recording. A second 25-minute recording covers the AWS-native stack. Together they are the two-version demonstration. This file is Databricks only.")
    mixed(doc, [
        ("Spoken lines ", {"size": 11}),
        ("are in roman type. Read them. Do not paraphrase Element labels or the five strategic-prompt answers.", {"size": 11}),
    ])
    mixed(doc, [
        ("Clicks ", {"size": 11}),
        ("are in gold ", {"size": 11}),
        ("[DO THIS]", {"bold": True, "italic": True, "size": 11, "color": GOLD}),
        (" boxes. Do the click, then keep reading. Do not stare at a spinner in silence.", {"size": 11}),
    ])
    speak(doc, "Clock marks are the latest you should still be on that beat. If you are late, jump to the Catch-up table. Pace is about 140 words per minute.")
    speak(doc, "Forward only. One Compass window after you enter the app. Do not Restore baseline. Do not retrain 04 or 04b. Do not unpause file-arrival. Do not invent an MFA form or a Deploy button inside Compass.")

    heading(doc, "25-minute clock")
    table(
        doc,
        ["Clock", "Where", "Live action", "Must say"],
        [
            ["0:00–2:00", "BEFORE Compass\nIdP / workspace login", "Live MFA at the identity provider", "Element 1  ·  Zero Trust  ·  least privilege"],
            ["2:00–4:00", "BEFORE Compass\nGitHub + terminal", "Show databricks.yml  ·  databricks bundle deploy -t poc", "Element 2  ·  most of IaC / CI-CD"],
            ["4:00–4:40", "Compass → Infrastructure", "Point at Estate tiles", "Element 2 remainder  ·  inventory of that deploy"],
            ["4:40–10:30", "Ingestion", "Ingest selected files  →  Hold  →  Start stream", "Element 3  ·  (a)  ·  (d)  ·  400→408"],
            ["10:30–13:00", "Catalog", "Open lineage  ·  Registry  ·  Quality  ·  tags", "Element 4  ·  (e)"],
            ["13:00–18:00", "Analytics", "Score registered models  ·  Scores  ·  Drift", "Element 5  ·  (b)"],
            ["18:00–20:30", "Portfolio", "Search quantum  ·  Accept/Defer  ·  daily brief", "Element 6"],
            ["20:30–25:00", "Export", "Execute export  ·  Statement API", "Element 7  ·  (c)"],
        ],
        col_widths=[1.15, 1.7, 2.15, 1.9],
    )
    mixed(doc, [
        ("Never cut: ", {"bold": True, "size": 11}),
        ("live MFA before the app, live repo/bundle before the app, Estate, 400→408, Hold, stream heartbeat, Catalog Explorer lineage, Score, search, Statement API, all five prompts.", {"size": 11}),
    ])
    mixed(doc, [
        ("Cut first if late: ", {"bold": True, "size": 11}),
        ("Quality tab, Policies tab, Anomalies or Forecast (keep one), Budget tab, download.", {"size": 11}),
    ])

    heading(doc, "Prompt map — say the letter out loud")
    table(
        doc,
        ["Prompt", "When", "One-line answer on this tape"],
        [
            ["(a) Legacy footprint", "During Ingest", "Strangler-fig. Same bronze. Volume landing. Rollback = delete-the-batch."],
            ["(b) Financial / budget", "During Score", "ERP + RF / IsolationForest / OLS. Protect ON_TARGET. Move $ off AT_RISK + TREND-DECLINE."],
            ["(c) Zero Trust / IL5", "On Export", "Three planes. App SP ≠ you. Gold SELECT. This cell is commercial AWS Moderate, not IL5."],
            ["(d) DR / RTO / RPO", "During Start stream", "Time travel. Volume durability. Serverless warehouse. Do not restore on camera."],
            ["(e) Vendor / lifecycle", "On Catalog tags", "vendor / license_id / renewal_date. Quality SLO. Last-good gold stays."],
        ],
        col_widths=[1.6, 1.5, 3.8],
    )

    heading(doc, "Night before — not recorded")
    bullets = [
        "Workspace Git folder: git pull main. Redeploy / restart onr-demo-poc.",
        "Start onr demo warehouse, onr demo cluster (your 04 / 04b), and onr demo ml (app Score). Leave them running.",
        "Compass Home = live 400, not fixture. If fixture: sql/grant_app_principal.sql + warehouse CAN USE for the app service principal.",
        "If silver ≠ 400: run 05_reset_demo.py on onr demo cluster.",
        "Run all 04_mlflow_grant_model.py, then 04b_funding_anomaly.py. Champion alias on funding_anomaly_detector. Do not train on camera.",
        "onr demo ml: Dedicated to the app SP, mlflow library installed, app SP has CAN RESTART + CAN ATTACH TO.",
        "Volume /Volumes/onr_demo/bronze/landing/_staged/batch_live_grants.csv exists.",
        "Mute notifications. 1920×1080, zoom 110–125%, hide bookmarks.",
        "Park the IdP login page, the GitHub repo (databricks.yml), and a terminal in the Databricks CLI folder. Compass is closed until 4:00.",
    ]
    for b in bullets:
        para = doc.add_paragraph(style="List Number")
        para.clear()
        r = para.add_run(b)
        _set_run(r, size=11)
        para.paragraph_format.space_after = Pt(3)

    heading(doc, "What is open")
    table(
        doc,
        ["Window", "Parked on", "When it appears"],
        [
            ["IdP / SSO", "Workspace login", "0:00"],
            ["GitHub", "jstrick9/onr_demo  ·  databricks.yml", "2:00"],
            ["Terminal", "Repo root, CLI authenticated", "2:30"],
            ["Compass  ·  onr-demo-poc", "Infrastructure", "4:00 — not before"],
        ],
        col_widths=[1.7, 3.2, 2.0],
    )

    # ---------- BEFORE APP ----------
    heading(doc, "Part A — before Compass  (0:00–4:00)")
    p(doc, "Element 1 and most of Element 2 happen outside this app. The spoken lines below are the Databricks version of that cold open so this 25-minute file still covers official 11.3. If your live MFA and live deploy are already in the AWS-native recording, still do a short Databricks-specific pass here — different IdP session, different repo, different command.", size=11, space_after=10)

    heading(doc, "0:00–2:00  ·  Element 1  ·  Secure access, MFA, Zero Trust", 2)
    do_this(doc, "Browser on the Databricks workspace login — not Compass. Start recording. Complete the identity-provider sign-in including the MFA challenge. Do not skip the challenge. Do not screenshot it later.")
    speak(doc, "This is the Databricks path of the ONR Code 08 demonstration. Everything you will see is mock, synthetic Compass data. No CUI, no PII, no classified.")
    speak(doc, "Element one: secure access, authentication, and Zero Trust.")
    speak(doc, "I am authenticating at the workspace identity provider — not inside an application form. Multi-factor authentication is the IdP challenge you just saw. There will be no password field in the Compass console. That is the Zero Trust control plane: identity is evaluated before the data plane opens.")
    speak(doc, "Least privilege starts here. The human session is mine. The Databricks App we will open has its own service principal. It does not borrow this token. Analysts SELECT gold. They never see bronze. Continuous authorization is the audit we will write on every search and every export.")
    speak(doc, "This cell is unclassified mock on commercial AWS — FedRAMP Moderate. The IL5 production cell is GovCloud, PrivateLink, customer-managed KMS. I will return to that on Element seven. We are not claiming this POC is IL5.")
    do_this(doc, "Leave the workspace signed in. Do not open Compass yet.")

    heading(doc, "2:00–4:00  ·  Element 2  ·  IaC, version control, CI/CD  (most of it)", 2)
    do_this(doc, "Switch to GitHub — github.com/jstrick9/onr_demo — file databricks.yml. Scroll enough to show volumes, the app, the paused file-arrival job, and the SDP pipeline. Do not scroll the whole file.")
    speak(doc, "Element two: infrastructure as code and automation.")
    speak(doc, "The demonstration estate is a Databricks Asset Bundle in this repository. The file is databricks.yml. Target poc. It manages the landing Volume, the checkpoint Volume, the Databricks App onr-demo-poc, the paused file-arrival job, and the SDP pipeline. It does not create the SQL warehouse or the clusters — those are pre-existing named compute: onr demo warehouse, onr demo cluster, onr demo ml.")
    do_this(doc, "Terminal at the repo root. Show the command. If you have already deployed tonight, do not apply again — read the command and the last successful deploy. If you have not deployed, run it and keep talking.")
    speak(doc, "Deploy is this command: databricks bundle deploy dash t poc. Version-controlled in Git. The same file a GitHub Action would call. There is no Deploy button in the application we are about to open. Configuration is reviewed in pull requests, not clicked in a console.")
    speak(doc, "That is the bulk of Element two. Compass will show the inventory that this deploy produced — not a second provision.")
    do_this(doc, "Open the Databricks App onr-demo-poc. Land on Infrastructure. Sidebar order is Home, Infrastructure, Ingestion, Catalog, Analytics, Portfolio, Export. Do not start on Home unless Active grants is not 400 — if it says fixture, stop and do not fake ingest.")

    # ---------- APP ----------
    heading(doc, "Part B — Compass  (4:00–25:00)")
    p(doc, "Partial Element 2, then Elements 3 through 7, with strategic prompts (a) through (e) woven in.", italic=True, color=MUTED, space_after=10)

    heading(doc, "4:00–4:40  ·  Element 2 remainder  ·  Infrastructure", 2)
    do_this(doc, "Cursor on Estate. Three tiles: warehouse, score cluster, file-arrival job. Do not open Full bundle unless a reviewer asks. Do not look for a Deploy button — there is not one.")
    speak(doc, "This is the inventory of the bundle we just walked. Warehouse onr demo warehouse — serverless SQL, serving this console. Score cluster onr demo ml — dedicated to the app service principal. File-arrival job — paused. The app you are looking at is onr-demo-poc. Catalog onr_demo, medallion bronze, silver, gold, and app.")
    speak(doc, "Logged-in user in the sidebar is my IdP session. The app is a different principal. That is the same Zero Trust split I signed in with.")
    do_this(doc, "Click Ingestion. Do not return to Infrastructure.")

    heading(doc, "4:40–10:30  ·  Element 3  ·  Ingestion  ·  prompts (a) and (d)", 2)
    speak(doc, "Element three: automated ingestion, data operations, and streaming.")
    speak(doc, "A new grants file has arrived. I am not recoding a pipeline. The staged files are already on the Unity Catalog Volume. Inbound grants is the good file. Quarantine sample is three bad rows — empty grant number, negative amount, and a duplicate.")
    do_this(doc, "Confirm both Inbound grants and Quarantine sample are selected. Click Ingest selected files. Keep talking.")
    speak(doc, "Strategic prompt (a) — sustainment of the legacy footprint.")
    speak(doc, "We do not cut over the D-and-A Portal, the reporting stack, or the existing ETL in a weekend. The pattern is strangler-fig coexistence. New files land in /Volumes/onr_demo/bronze/landing/grants/ — governed object storage, not a DBFS mount. The warehouse path this button just used, and the Auto Loader path I will start next, write the same bronze Delta table. Legacy reports keep reading gold over JDBC. When a report is ready to retire, we point it at the table this app already uses. Rollback is delete-the-batch, not rewrite-the-estate.")
    do_this(doc, "Point at Active grants 400 → 408 and Quarantined +3. Point at the Hold tray — chips empty, dup, amt.")
    speak(doc, "There it is. Eight rows published. The three quarantine rows never entered bronze — they are in app.quarantine_log. One live grant published with a warning, missing abstract, in app.quality_findings. Quality is the tab if a reviewer wants the scoreboard. We do not need it on this clock.")
    speak(doc, "Same Element three. That button was the warehouse SQL path. Next is file-arrival. I am not leaving this console.")
    do_this(doc, "Click Start stream. Do not Restore baseline — that control is at the bottom of the page and is not on this tape. If bronze does not tick, Workspace strip 01b stream, Run all, come straight back.")
    speak(doc, "The console just wrote a new batch_live_grants_stream file onto the landing Volume and loaded it into the same bronze table. That is the stream proof: file arrival, not a scheduled batch. Auto Loader — cloudFiles — is the Spark job on that Volume. Open 01b if a reviewer wants the job run. We do not wait on Jobs serverless cold start on this tape.")
    do_this(doc, "Point at bronze count, last 2 min, last file ago, then Delta time travel.")
    speak(doc, "Bronze ticks. Silver stays at four hundred and eight because silver dedupes on grant_no. Kafka or Kinesis would be the equivalent bus on another estate. Here the open equivalent is a Unity Catalog Volume plus Auto Loader.")
    speak(doc, "Strategic prompt (d) — disaster recovery, resilience, and failover.")
    speak(doc, "Contract targets: RPO fifteen minutes for gold, essentially zero for bronze landing — the file is still sitting in the Volume. Delta time-travel is the row-level RPO — baseline snapshot versus now. We are not restoring on camera. RTO thirty minutes to serving gold, about five minutes to serving the app. The warehouse is serverless. The app is already deployed. The all-purpose cluster is not on the serving path. Annual DR is non-disruptive: pause file-arrival, restore yesterday’s gold into a onr_demo_dr catalog, point a clone of the app at it, validate, tear down.")
    do_this(doc, "Click Catalog. Do not return to Ingestion.")

    heading(doc, "10:30–13:00  ·  Element 4  ·  Catalog  ·  prompt (e)", 2)
    speak(doc, "Element four: data governance, quality, and cataloging.")
    do_this(doc, "Click Open lineage · gold.grants_summary. That is the only workspace jump. If the graph is empty, open gold.grants_summary from the same explorer. Come back immediately.")
    speak(doc, "This is Unity Catalog’s native lineage — landing Volume, to bronze, to silver, to gold, to the app. That graph is the Element four visual. This console is the operator surface, not the system of record.")
    do_this(doc, "Back on Catalog — Registry. Show bronze, silver, gold, app.")
    speak(doc, "Four schemas. The eight grants we just ingested are already registered. Metadata — source file, ingest time, tags — is on the table, not in a spreadsheet.")
    do_this(doc, "Quality tab. Point at the health scores. Do not open every expander.")
    speak(doc, "Completeness, accuracy, consistency, timeliness. A lapsed vendor feed shows up here as a timeliness drop, not a blank dashboard.")
    do_this(doc, "Policies and tags. Point at data_source = mock.")
    speak(doc, "Strategic prompt (e) — data vendor and lifecycle management.")
    speak(doc, "Every external feed is a licensed product: owner, renewal date, quality SLO. Today the tags are data_source, domain, data_sensitivity. In production we add vendor, license_id, renewal_date. Usage is metered in app.export_history and app.search_history. If a subscription stops, Auto Loader has nothing new. Last-good gold stays. We do not auto-delete.")
    do_this(doc, "Click Analytics. Do not return to Catalog.")

    heading(doc, "13:00–18:00  ·  Element 5  ·  Analytics  ·  prompt (b)", 2)
    speak(doc, "Element five: decision-support analytics and modeling.")
    speak(doc, "The models were trained last night and registered in Unity Catalog. I am not retraining on camera. I am triggering a live score against the portfolio we just ingested — including the eight new grants.")
    do_this(doc, "Click Score registered models. Point at Scores — Fund, Review, Defer, Flagged. Do not hunt tabs first. If the run does not submit, Workspace strip 04c score, Run all, come straight back. Keep talking.")
    speak(doc, "Strategic prompt (b) — financial and budgetary analytical integration.")
    speak(doc, "Financial execution is a first-class feed: twelve hundred ERP lines — budget, actual, execution rate — into gold.budget_execution. Three complementary models, all on this ingested portfolio, not a second dataset.")
    speak(doc, "Descriptive: the Portfolio page a Code 08 resource officer opens — dollars, execution, ON_TARGET, WARNING, AT_RISK.")
    speak(doc, "Predictive: a Random Forest large-award classifier — Fund, Review, or Defer — registered as onr_demo.gold.grant_large_award. An IsolationForest for budget spike, execution collapse, and low-return concentration — funding_anomaly_detector at alias champion. And ordinary least squares, ols_fy_v1, two-year horizon, ninety-five percent band, trend IDs TREND-ACCEL, TREND-STEADY, TREND-DECLINE. That is OLS, not Prophet.")
    speak(doc, "Prescriptive: protect ON_TARGET, move dollars off AT_RISK and TREND-DECLINE, review large-award concentration. Engineer owns bronze and silver. Scientist owns the registered models. Analyst owns gold and the daily brief. Same catalog.")
    do_this(doc, "Point at Resource action. Then Drift — program-mix PSI, award-size PSI, Fund share.")
    speak(doc, "That is feature and score mix versus the baseline snapshot, not a fake accuracy drop. The live grants moved the mix.")
    do_this(doc, "Predictions tab. Point at model_name. Then Anomalies. Then Forecasting — orange horizon and one TREND-DECLINE row. Skip Metrics.")
    speak(doc, "The Resource action sentence is the close a resource officer would sign — Defer dollars off one area onto AT_RISK plus TREND-DECLINE. The declining program is the reallocation candidate on the next page.")
    do_this(doc, "Click Portfolio. Do not return to Analytics.")

    heading(doc, "18:00–20:30  ·  Element 6  ·  Portfolio", 2)
    speak(doc, "Element six: unified dashboard, visualizations, and process automation.")
    speak(doc, "A non-technical leader does not need SQL. Active grants is still four hundred and eight.")
    do_this(doc, "Search box on this page. Type quantum. Press enter. Point at Routing. Accept or Defer one flagged grant.")
    speak(doc, "Search is live against gold. It is written to app.search_history. That is the audit Zero Trust asked for, and it is the usage meter for prompt (e).")
    do_this(doc, "Click Generate daily brief. Wait for the letterhead. Stay on this page.")
    speak(doc, "Automated summary. Classification banner, three bullets, one recommended action. If Foundation Model serving is on, this is ai_query. If not, it is the structured template. Either way a row lands in app.daily_briefs. That is process automation — not a staffer writing the morning book.")
    do_this(doc, "Budget tab. Point at one AT_RISK row. Then click Export.")
    speak(doc, "Same gold the forecast used. AT_RISK plus TREND-DECLINE is the reallocation set. Extract is next.")

    heading(doc, "20:30–25:00  ·  Element 7  ·  Export  ·  prompt (c)", 2)
    speak(doc, "Element seven: interoperability, data portability, and secure export.")
    do_this(doc, "Date range is already 2025 to 2026. Leave CSV and Parquet on. Dataset Grants Summary. Click Execute export. Download Parquet or CSV once. Stay on this screen.")
    speak(doc, "Filtered bulk extract — not SELECT star. Open formats: CSV, JSON, Parquet. Schema travels with the file: grant_no, program_area, amount_usd, awardee.")
    do_this(doc, "Open History. Point at the new row.")
    speak(doc, "app.export_history — who, what, filter, row count. Continuous authorization, not a static password file.")
    do_this(doc, "Click Execute live Statement API call. Point at the statement receipt — statement_id, SUCCEEDED, row_count, warehouse, elapsed.")
    speak(doc, "This is the live, documented Databricks Statement Execution REST API — POST /api/2.0/sql/statements — OAuth, short-lived token, same warehouse this dashboard uses. That is what Advana or Cloud One would call. It is not a Databricks-only extract and it is not a fictional host.")
    speak(doc, "Strategic prompt (c) — Zero Trust and IL5.")
    speak(doc, "Three planes, same identity. The app has its own service principal; it does not borrow mine. The warehouse and the cluster are separate compute — micro-segmentation at the control plane. Unity Catalog is the data-plane firewall. Least privilege: analysts SELECT gold, they never see bronze. This cell is unclassified mock on commercial AWS — FedRAMP Moderate. The IL5 production cell is GovCloud, PrivateLink, customer-managed KMS, continuous compliance against the same least-privilege grants. We are not claiming this POC is IL5.")
    speak(doc, "Tomorrow another CSV lands on the same Volume. Same gold. Same registered models — we rescore from this console. Mock data only.")
    do_this(doc, "Stop talking. Leave Export on screen.")

    heading(doc, "Catch-up  ·  if the clock is ahead of you")
    table(
        doc,
        ["If the clock says", "Skip and go to"],
        [
            ["2:30 and MFA is still spinning", "Finish the challenge. Do not skip MFA. Steal time from Catalog Policies later."],
            ["3:50 still in the repo", "Do not run deploy again. Say the command. Open Compass on Infrastructure at 4:00."],
            ["6:30 and Ingest is still spinning", "Keep talking (a). Do not open Quality."],
            ["9:30 and bronze has not ticked", "Say the file is on the Volume. Open 01b if needed. At 10:30 go to Catalog anyway."],
            ["12:30 still on Catalog", "Skip Policies. Say (e) in two sentences. Go to Analytics."],
            ["16:30 and Score is still running", "Stay on Analytics. Skip Metrics. Portfolio at 18:00."],
            ["19:30 still on Analytics", "Skip Anomalies or Forecast — keep one. Go to Portfolio search."],
            ["21:30 still on Portfolio", "Skip Budget. Go to Export."],
            ["24:00 and no API yet", "Skip download. Execute live Statement API immediately. Say (c) in four sentences."],
        ],
        col_widths=[2.6, 4.3],
    )

    heading(doc, "If something is down  ·  one sentence, then move")
    table(
        doc,
        ["Symptom", "Say, then do"],
        [
            ["Compass says fixture", "Warehouse or app service-principal grant is cold. Do not fake Ingest."],
            ["Ingest errors", "I will not debug IAM on camera. Skip to Catalog Open lineage."],
            ["Start stream fails", "Open stream notebook. Same Volume file. Come back."],
            ["Bronze does not tick", "File must sit under landing/grants, not _staged. Then Catalog."],
            ["Lineage graph empty", "Same graph lives on gold.grants_summary. Do not draw it in Streamlit."],
            ["Score fails", "Open scoring notebook. I will not train on camera. Last night’s gold tables."],
            ["Brief is the template", "Foundation Models are off. The structured brief still writes app.daily_briefs."],
            ["Statement API fails", "Same SQL on the warehouse cursor — still live. The curl is the Advana contract."],
            ["Widget clears on rerun", "Re-select. Do not Restore baseline."],
        ],
        col_widths=[2.2, 4.7],
    )

    heading(doc, "What you will not say")
    nos = [
        "This POC is IL5 or FedRAMP High.",
        "The forecast is Prophet or a neural net. It is OLS.",
        "The stream is Amazon Kinesis. It is a Unity Catalog Volume plus Auto Loader.",
        "Lineage is drawn inside Streamlit. Native lineage is Catalog Explorer.",
        "We trained the models just now. You scored from the registry. Training was last night.",
        "Drift is accuracy decay. It is feature and score mix versus the baseline snapshot.",
        "We rewrite the legacy estate. Strangler-fig only.",
        "MFA is this Compass form. MFA is the IdP — you already did it.",
        "I am deploying from this button. There is no deploy button. Deploy was databricks bundle deploy -t poc.",
        "Any second pass through Home. You entered on Infrastructure.",
    ]
    for n in nos:
        para = doc.add_paragraph(style="List Bullet")
        para.clear()
        r = para.add_run(n)
        _set_run(r, size=11)
        para.paragraph_format.space_after = Pt(2)

    heading(doc, "Close")
    speak(doc, "You have executed all seven scenario elements and narrated all five strategic prompts. Element 1 and most of Element 2 were live before Compass. Compass carried the Element 2 inventory and Elements 3 through 7. Stop on Export. Do not keep talking.")

    p(doc, "UNCLASSIFIED  //  MOCK DATA", bold=True, size=10, color=GOLD, space_before=18, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(OUT)
    print("wrote", OUT, "bytes", OUT.stat().st_size)


if __name__ == "__main__":
    build()
